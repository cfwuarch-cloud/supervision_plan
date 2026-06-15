# -*- coding: utf-8 -*-
"""
監造組織及權責分工 (Ch2) 轉換工具 v1.0
========================================
產出監造計畫第二章 docx，含 SVG 組織圖及權責分工表。

修正歷程：
  v1.0  2026/06/15  初始版本

作者：OpenCode Assistant / cfwuarch
版本：v1.0
最後更新：2026/06/15

相依套件：
  - python-docx>=1.0.0,<2.0.0
  - lxml>=4.9.0,<6.0.0

使用方法：
  python -X utf8 tables/ch2/convert_2.py

功能說明：
  1. 讀取 ch2_內容.md 文字內容
  2. 生成 SVG 組織圖（側輸出 svg 檔）
  3. 產出監造組織章節 docx

參數說明：
  -c, --content    文字內容 md 路徑（預設：./ch2_內容.md）
  -o, --output     輸出 docx 路徑
  --svg-output     SVG 組織圖輸出路徑
"""
import argparse
import os
import sys
import re
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))))
from common.docx_table import add_cell


def clean_text(text):
    text = text.replace('\n', '').replace('\r', '')
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[\s]+([、，,。．])', r'\1', text)
    return text.strip()


def read_content_md(path, placeholders=None):
    sections = {}
    if not path or not os.path.isfile(path):
        return sections
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    if placeholders:
        for k, v in placeholders.items():
            text = text.replace('{{' + k + '}}', v)
    current_heading = None
    current_lines = []
    for line in text.split('\n'):
        if line.startswith('# ') or line.startswith('## '):
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_lines).strip()
            current_heading = re.sub(r'^#+ ', '', line).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_lines).strip()
    return sections


def gen_org_chart_svg(svg_path):
    """生成監造組織圖 SVG（14px 底部間距規範）"""
    BOX_W = 180
    BOX_H = 50
    FONT_SIZE = 13
    LINE_H = 18
    GAP_X = 20
    GAP_Y = 50
    PAD_TOP = 20
    PAD_LEFT = 20

    org = [
        ('監造主持人', 0, 0),
        ('監造主任', 0, 1),
        ('土建監造\n工程師', -1, 2),
        ('機電監造\n工程師', 0, 2),
        ('品管人員', 1, 2),
        ('行政助理', 2, 2),
    ]

    def box_center(lvl, idx, total_lvl2=4):
        cx = PAD_LEFT + (total_lvl2 - 1) * (BOX_W + GAP_X) // 2 + idx * (BOX_W + GAP_X)
        if lvl == 0:
            cx = PAD_LEFT + (total_lvl2 - 1) * (BOX_W + GAP_X) // 2
        elif lvl == 1:
            cx = PAD_LEFT + (total_lvl2 - 1) * (BOX_W + GAP_X) // 2
        return cx

    svg_w = PAD_LEFT * 2 + 4 * (BOX_W + GAP_X)
    svg_h = PAD_TOP * 2 + 3 * (BOX_H + GAP_Y)

    nsmap = {None: 'http://www.w3.org/2000/svg'}
    svg = etree.Element('svg', nsmap=nsmap)
    svg.set('width', str(svg_w))
    svg.set('height', str(svg_h))
    svg.set('viewBox', f'0 0 {svg_w} {svg_h}')

    # 定義箭頭標記
    defs = etree.SubElement(svg, 'defs')
    marker = etree.SubElement(defs, 'marker')
    marker.set('id', 'arrow')
    marker.set('markerWidth', '10')
    marker.set('markerHeight', '10')
    marker.set('refX', '10')
    marker.set('refY', '5')
    marker.set('orient', 'auto')
    path_m = etree.SubElement(marker, 'path')
    path_m.set('d', 'M0,0 L10,5 L0,10 Z')
    path_m.set('fill', '#000')

    positions = {}
    for role, offset, lvl in org:
        if lvl == 0:
            x = PAD_LEFT + 2 * (BOX_W + GAP_X)
        elif lvl == 1:
            x = PAD_LEFT + 2 * (BOX_W + GAP_X)
        else:
            base_x = PAD_LEFT + (BOX_W + GAP_X)
            x = base_x + (offset + 1) * (BOX_W + GAP_X)
        y = PAD_TOP + lvl * (BOX_H + GAP_Y)
        positions[role] = (x, y, lvl)

    # 畫框 + 連線
    parent_map = {
        '監造主任': '監造主持人',
        '土建監造\n工程師': '監造主任',
        '機電監造\n工程師': '監造主任',
        '品管人員': '監造主任',
        '行政助理': '監造主任',
    }

    for role, (x, y, lvl) in positions.items():
        # 連線
        if role in parent_map:
            px, py, _ = positions[parent_map[role]]
            line = etree.SubElement(svg, 'line')
            line.set('x1', str(px + BOX_W // 2))
            line.set('y1', str(py + BOX_H))
            line.set('x2', str(x + BOX_W // 2))
            line.set('y2', str(y))
            line.set('stroke', '#000')
            line.set('stroke-width', '1.5')
            line.set('marker-end', 'url(#arrow)')

        # 矩形框
        rect = etree.SubElement(svg, 'rect')
        rect.set('x', str(x))
        rect.set('y', str(y))
        rect.set('width', str(BOX_W))
        rect.set('height', str(BOX_H))
        rect.set('fill', '#fff')
        rect.set('stroke', '#000')
        rect.set('stroke-width', '1.5')

        # 文字（支援多行）
        lines = role.split('\n')
        for li, line_txt in enumerate(lines):
            txt = etree.SubElement(svg, 'text')
            txt.set('x', str(x + BOX_W // 2))
            txt.set('y', str(y + BOX_H // 2 - (len(lines) - 1) * LINE_H // 2 + li * LINE_H))
            txt.set('text-anchor', 'middle')
            txt.set('font-size', str(FONT_SIZE))
            txt.set('font-family', '標楷體')
            txt.text = line_txt

    tree = etree.ElementTree(svg)
    tree.write(svg_path, encoding='utf-8', xml_declaration=True)
    print(f'  SVG 組織圖 → {os.path.basename(svg_path)}')
    return svg_path


def build_doc(content_path, output_path, svg_output_path):
    placeholders = {
        '土建監造人數': '2',
        '機電監造人數': '1',
        '品管人數': '1',
        '行政人數': '1',
    }
    sections = read_content_md(content_path, placeholders)

    # 生成 SVG
    gen_org_chart_svg(svg_output_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    doc.add_heading('第二章  監造組織及權責分工', level=1)

    # ——— 1. 監造組織 ———
    doc.add_heading('1  監造組織', level=2)

    # 1.1 組織架構
    doc.add_heading('(1) 架構', level=3)
    doc.add_paragraph('監造組織架構如下圖所示，含監造單位管理階層、工地部門及派駐人員職稱配置。')
    doc.add_paragraph()
    doc.add_paragraph(f'（SVG 組織圖：{os.path.basename(svg_output_path)}）')
    doc.add_paragraph()

    # 1.2 人員配置
    doc.add_heading('(2) 人員配置', level=3)
    doc.add_paragraph('依工程規模及契約、「公共工程施工品質管理作業要點」之規定，'
                      '檢討預定配置符合規定之工地人員人數如下：')

    staff = [
        ('監造主持人', '1 名', '綜理監造業務'),
        ('監造主任', '1 名', '工地現場監造事務總責'),
        ('土建監造工程師', placeholders['土建監造人數'] + ' 名', '土建工程抽查及材料抽驗'),
        ('機電監造工程師', placeholders['機電監造人數'] + ' 名', '機電工程抽查及設備測試'),
        ('品管人員', placeholders['品管人數'] + ' 名', '品質文件審查及稽核'),
        ('行政助理', placeholders['行政人數'] + ' 名', '文書行政事務'),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['職稱', '人數', '工作重點']):
        hdr[i].text = txt
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for role, num, desc in staff:
        row = table.add_row().cells
        row[0].text = role
        row[1].text = num
        row[2].text = desc

    # ——— 2. 工作職掌 ———
    doc.add_heading('2  工作職掌', level=2)
    doc.add_paragraph('依服務契約規定，明確劃分所有監造組織內職稱人員應辦理之工作內容及重點。')

    duties = [
        ('監造主持人', '1. 綜理監造業務之規劃、督導與管理\n'
                       '2. 重大品質問題之決策\n'
                       '3. 督導監造計畫之執行與修正'),
        ('監造主任', '1. 工地現場監造事務之總負責\n'
                     '2. 審查廠商施工計畫及品質計畫\n'
                     '3. 主持各項協調會議\n'
                     '4. 督導監造人員執行抽查及抽驗作業'),
        ('土建監造工程師', '1. 土建工程之施工抽查\n'
                          '2. 材料及設備之抽驗\n'
                          '3. 施工查驗紀錄之製作\n'
                          '4. 不合格品缺失之追蹤改善'),
        ('機電監造工程師', '1. 機電工程之施工抽查\n'
                          '2. 設備功能運轉測試之抽驗\n'
                          '3. 機電系統整合之查證\n'
                          '4. 機電相關文件審查'),
        ('品管人員', '1. 品質相關文件之審查\n'
                     '2. 品質稽核之規劃與執行\n'
                     '3. 不合格品追蹤管制\n'
                     '4. 品質數據之整理分析'),
        ('行政助理', '1. 文件之收發、登錄與歸檔\n'
                     '2. 會議通知及紀錄\n'
                     '3. 行政事務及後勤支援'),
    ]

    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text = '職稱'
    hdr2[1].text = '工作職掌'
    for role, duties_text in duties:
        row = table2.add_row().cells
        row[0].text = role
        row[1].text = duties_text

    # ——— 3. 開工前協調會議 ———
    doc.add_heading('3  開工前協調會議', level=2)
    doc.add_paragraph('工程決標後開工前及各分項工程施工前，應召開開工前協調會議，'
                      '宣達權責分工，將工程設計理念、監造標準、施工規範及契約重要規定，'
                      '正確有效地傳遞予施工廠商。')

    doc.save(output_path)
    print(f'Ch2 監造組織及權責分工 → {os.path.basename(output_path)}')


def main():
    parser = argparse.ArgumentParser(description='Ch2 監造組織及權責分工轉換工具 v1.0')
    parser.add_argument('-c', '--content', default='./ch2_內容.md',
                        help='文字內容 md 路徑')
    parser.add_argument('-o', '--output', default='../../output/Ch2_監造組織及權責分工.docx',
                        help='輸出 docx 路徑')
    parser.add_argument('--svg-output', default='../../output/Ch2_組織圖.svg',
                        help='SVG 組織圖輸出路徑')
    args = parser.parse_args()

    base = os.path.dirname(os.path.abspath(__file__))
    content_path = os.path.normpath(os.path.join(base, args.content))
    output_path = os.path.normpath(os.path.join(base, args.output))
    svg_output_path = os.path.normpath(os.path.join(base, args.svg_output))

    build_doc(content_path, output_path, svg_output_path)


if __name__ == '__main__':
    main()
