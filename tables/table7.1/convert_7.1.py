#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
圖7.1 輕隔間施工抽查流程圖 — 產製腳本

修正歷程：
  v1.0  2026-06-14  OpenCode Assistant  初版

作者：OpenCode Assistant / cfwuarch
版本：v1.0
最後更新：2026-06-14

相依套件：
  python-docx>=1.1.0,<2.0.0
  lxml>=4.9.0,<6.0.0

使用方法：
  python -X utf8 tables/table7.1/convert_7.1.py

功能說明：
  從 data/圖7.1_輕隔間施工抽查流程圖.docx 提取 VML 流程圖元件，
  產出相同內容的 Word 文件至 output/圖7.1_輕隔間施工抽查流程圖.docx。

參數說明：
  無（固定產自 data/ 目錄下的原始檔）
"""

import os, sys, copy
from lxml import etree

from docx import Document
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# 路徑
# ---------------------------------------------------------------------------
BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
SRC = os.path.join(BASE, "data", "圖7.1_輕隔間施工抽查流程圖.docx")
OUT = os.path.join(BASE, "output", "圖7.1_輕隔間施工抽查流程圖.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

# ---------------------------------------------------------------------------
# 1. 從原始檔提取 VML group XML 及 styles
# ---------------------------------------------------------------------------
def extract_vml_group(src):
    """從原始 docx 取出 VML group 的 lxml 元素"""
    doc = Document(src)
    body = doc.element.body

    # VML 放在 <w:pict> 內 → 在 w:r 內
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    pict = body.find(f".//{{{ns_w}}}pict")
    if pict is None:
        raise RuntimeError("找不到 VML pict 元素")
    # pict 內第一個子元素就是 v:group
    vml_group = list(pict)[0]
    # deepcopy 以便獨立使用
    return copy.deepcopy(vml_group)


# ---------------------------------------------------------------------------
# 2. 建構新文件
# ---------------------------------------------------------------------------
def build_document(vml_group):
    doc = Document()

    # ── 頁面設定 ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Emu(11906 * 914400 // 1440)   # 11906 twip → EMU
    section.page_height = Emu(16838 * 914400 // 1440)
    section.top_margin = Cm(2.54)       # 1440 twip ≈ 2.54cm
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.175)     # 1800 twip ≈ 3.175cm
    section.right_margin = Cm(3.175)

    # ── P0: 標題 ──────────────────────────────────────────────────────────
    p0 = doc.add_paragraph()
    run0 = p0.add_run("1.輕隔間施工工作")
    run0.bold = True
    run0.font.size = Pt(14)

    # ── P1: 說明文字 ──────────────────────────────────────────────────────
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = Pt(12)  # atLeast 228600 EMU = 12pt
    p1.paragraph_format.line_spacing_rule = 3  # WD_LINE_SPACING.AT_LEAST
    run1 = p1.add_run("輕隔間施工流程及檢驗程序      以〝＊〞標明施工檢驗停留點")
    run1.font.name = "標楷體"
    run1.font.size = Pt(11)
    # 設定東亞字型
    rPr = run1._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        rFonts.tail = None
    rFonts.set(qn("w:eastAsia"), "標楷體")

    # ── P2: VML 流程圖容器 ──────────────────────────────────────────────
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)

    # 建立 w:r / w:pict 並插入 VML group
    run2 = p2.add_run("")

    # 取得 pict 元素的字串序列化物件
    nsmap = {
        "v": "urn:schemas-microsoft-com:vml",
        "o": "urn:schemas-microsoft-com:office:office",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "w10": "urn:schemas-microsoft-com:office:word",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    pict_xml = etree.Element(qn("w:pict"), nsmap=nsmap)
    pict_xml.append(vml_group)

    # 塞入 run
    r_elem = run2._r
    # 清除 run 內原本的文字元素
    for child in list(r_elem):
        r_elem.remove(child)
    r_elem.append(pict_xml)

    # 也保留 text fallback（P2 在 Word 中顯示的文字備援）
    fallback_run = p2.add_run("封板骨架組立＊檢驗停留點＊檢驗停留點NO放樣填充玻璃棉YES門框安裝預留開口配管配線加強構件完成防火一小時時效隔間")

    # ── P3～P15: 空白段（保留版面空間） ─────────────────────────────────
    for _ in range(13):
        doc.add_paragraph()

    # ── P16: 圖說 ────────────────────────────────────────────────────────
    p16 = doc.add_paragraph()

    # ── P17: 結尾空白段 ────────────────────────────────────────────────
    doc.add_paragraph()
    run16a = p16.add_run("圖")
    run16b = p16.add_run("7-1 ")
    run16c = p16.add_run("輕隔間施工抽查流程圖")

    # ── 設定 sectPr 屬性 (對齊原始檔) ───────────────────────────────────
    sect = doc.sections[0]
    sect_pr = sect._sectPr

    cols = etree.SubElement(sect_pr, qn("w:cols"))
    cols.set(qn("w:space"), "425")

    doc_grid = etree.SubElement(sect_pr, qn("w:docGrid"))
    doc_grid.set(qn("w:linePitch"), "360")
    doc_grid.set(qn("w:type"), "lines")

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    print("從原始檔提取 VML 流程圖元件...")
    vml = extract_vml_group(SRC)

    print("建構新文件...")
    doc = build_document(vml)

    print(f"儲存至 {OUT}")
    doc.save(OUT)
    print("完成")


if __name__ == "__main__":
    main()
