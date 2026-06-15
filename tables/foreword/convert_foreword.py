# -*- coding: utf-8 -*-
"""
前言轉換工具 v1.0
===================
產出監造計畫「前言」docx。

修正歷程：
  v1.0  2026/06/15  初始版本

作者：OpenCode Assistant / cfwuarch
版本：v1.0
最後更新：2026/06/15

使用方法：
  python -X utf8 tables/foreword/convert_foreword.py

參數說明：
  -c, --content    文字內容 md 路徑
  -o, --output     輸出 docx 路徑
  --project-name   專案名稱
"""
import argparse
import os
import sys
import re
from docx import Document
from docx.shared import Pt, Cm

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))))


def read_content_md(path, placeholders=None):
    sections = {}
    if not path or not os.path.isfile(path):
        return sections
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    if placeholders:
        for k, v in placeholders.items():
            text = text.replace('{{' + k + '}}', v)
    current_heading = None
    current_lines = []
    for line in text.split('\n'):
        if line.startswith('# ') or line.startswith('## '):
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_lines).strip()
            current_heading = re.sub(r'^#+ ', '', line).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_lines).strip()
    return sections


def build_doc(content_path, output_path, project_name):
    sections = read_content_md(content_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    doc.add_heading(f'{project_name}', level=0)
    doc.add_heading('監造計畫', level=0)
    doc.add_paragraph()

    doc.add_heading('前  言', level=1)

    # ——— 編製依據 ———
    doc.add_heading('編製依據', level=2)
    doc.add_paragraph('本監造計畫依下列法規及契約規定編製：')
    for item in ['「公共工程施工品質管理作業要點」第八點',
                 '本工程服務契約',
                 '本工程工程契約（含規範及圖說）']:
        doc.add_paragraph(item, style='List Bullet')

    # ——— 計畫目的 ———
    doc.add_heading('計畫目的', level=2)
    doc.add_paragraph('公共工程三級品管制度的落實執行，攸關公共工程品質至鉅。'
                      '其中第二層級之品質保證扮演著工程品質把關之角色。')

    # ——— 三級品管制度 ———
    doc.add_heading('三級品管制度', level=2)
    tiers = [
        ('第一層級', '廠商施工品質管制系統（品質計畫）'),
        ('第二層級', '主辦機關及監造單位品質保證系統（監造計畫）'),
        ('第三層級', '工程會及主管機關工程施工查核機制'),
    ]
    for tier, desc in tiers:
        doc.add_paragraph(f'{tier}：{desc}')

    doc.add_paragraph()
    doc.add_paragraph('監造計畫屬於第二層級品質保證，為確保工程的施工成果能符合設計及規範，'
                      '主辦機關應建立施工品質保證系統，設立監造組織，訂定監造計畫，'
                      '並落實執行，以確保工程可如期如質完成。')

    # ——— 計畫注意事項 ———
    doc.add_heading('計畫注意事項', level=2)
    doc.add_paragraph('一、監造計畫應對人力規劃、監督作法、監督紀錄，'
                      '及就廠商之施工計畫、品質計畫等如何有效審查，作有系統之規劃。')
    doc.add_paragraph('二、監造單位應對於下列各項提出具體作法並紀錄其重點：')
    for item in ['查證廠商相關書面作業落實執行狀況',
                 '材料取樣、抽驗檢試驗及數據整理分析管制',
                 '對現場施工工法、施工管控、施工過程與結果作持續性監督與查證',
                 '不合格品瑕疵列管、改善追蹤管制',
                 '對廠商內部品質稽核結果及自主品管落實度做進一步之稽核']:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph('三、製作監造計畫時，除依契約及作業要點規定辦理外，'
                      '另應參酌其他法令規定。')
    doc.add_paragraph('四、監造計畫應於工程發包前提報甲方審核，'
                      '並於工程決標前完成核定程序。')

    doc.save(output_path)
    print(f'前言 → {os.path.basename(output_path)}')


def main():
    parser = argparse.ArgumentParser(description='前言轉換工具 v1.0')
    parser.add_argument('-c', '--content', default='./foreword_內容.md')
    parser.add_argument('-o', '--output', default='../../output/前言.docx')
    parser.add_argument('--project-name', default='（專案名稱）')
    args = parser.parse_args()
    base = os.path.dirname(os.path.abspath(__file__))
    build_doc(
        os.path.normpath(os.path.join(base, args.content)),
        os.path.normpath(os.path.join(base, args.output)),
        args.project_name)


if __name__ == '__main__':
    main()
