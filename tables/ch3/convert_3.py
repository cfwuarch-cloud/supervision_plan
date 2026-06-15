# -*- coding: utf-8 -*-
"""
品質計畫審查作業程序 (Ch3) 轉換工具 v1.0
==========================================
產出監造計畫第三章 docx，含審查流程圖 SVG 及表3.1 審查意見表。

修正歷程：
  v1.0  2026/06/15  初始版本

作者：OpenCode Assistant / cfwuarch
版本：v1.0
最後更新：2026/06/15

相依套件：
  - python-docx>=1.0.0,<2.0.0
  - lxml>=4.9.0,<6.0.0

使用方法：
  python -X utf8 tables/ch3/convert_3.py

功能說明：
  1. 讀取 ch3_內容.md 文字內容
  2. 生成審查流程 SVG
  3. 產出品審審查章節 docx

參數說明：
  -c, --content    文字內容 md 路徑
  -o, --output     輸出 docx 路徑
  --svg-output     SVG 流程圖輸出路徑
"""
import argparse
import os
import sys
import re
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))))
from common.docx_table import add_cell


def read_content_md(path, placeholders=None):
    sections = {}
    if not path or not os.path.isfile(path):
        return sections
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    if placeholders:
        for k, v in placeholders.items():
            text = text.replace('{{' + k + '}}', v)
    current_heading = None
    current_lines = []
    for line in text.split('\n'):
        if line.startswith('# ') or line.startswith('## '):
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_lines).strip()
            current_heading = re.sub(r'^#+ ', '', line).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_lines).strip()
    return sections


def gen_flow_chart_svg(svg_path):
    """生成品質計畫審查流程 SVG（14px 底部間距規範）"""
    BOX_W = 160
    BOX_H = 45
    DIA_W = 140
    DIA_H = 70
    FONT_SIZE = 13
    GAP_Y = 30
    PAD_TOP = 30
    PAD_LEFT = 50

    svg_w = 400
    svg_h = 500

    nsmap = {None: 'http://www.w3.org/2000/svg'}
    svg = etree.Element('svg', nsmap=nsmap)
    svg.set('width', str(svg_w))
    svg.set('height', str(svg_h))
    svg.set('viewBox', f'0 0 {svg_w} {svg_h}')

    defs = etree.SubElement(svg, 'defs')
    marker = etree.SubElement(defs, 'marker')
    marker.set('id', 'arrow')
    marker.set('markerWidth', '10')
    marker.set('markerHeight', '10')
    marker.set('refX', '10')
    marker.set('refY', '5')
    marker.set('orient', 'auto')
    pm = etree.SubElement(marker, 'path')
    pm.set('d', 'M0,0 L10,5 L0,10 Z')
    pm.set('fill', '#000')

    def box(x, y, text, w=BOX_W, h=BOX_H):
        r = etree.SubElement(svg, 'rect')
        r.set('x', str(x))
        r.set('y', str(y))
        r.set('width', str(w))
        r.set('height', str(h))
        r.set('fill', '#fff')
        r.set('stroke', '#000')
        r.set('stroke-width', '1.5')
        lines = text.split('\n')
        for li, t in enumerate(lines):
            tx = etree.SubElement(svg, 'text')
            tx.set('x', str(x + w // 2))
            tx.set('y', str(y + h // 2 - (len(lines)-1)*9 + li*18 + 5))
            tx.set('text-anchor', 'middle')
            tx.set('font-size', str(FONT_SIZE))
            tx.set('font-family', '標楷體')
            tx.text = t

    def diamond(x, y, text, w=DIA_W, h=DIA_H):
        p = etree.SubElement(svg, 'polygon')
        pts = f'{x+w//2},{y} {x+w},{y+h//2} {x+w//2},{y+h} {x},{y+h//2}'
        p.set('points', pts)
        p.set('fill', '#fff')
        p.set('stroke', '#000')
        p.set('stroke-width', '1.5')
        tx = etree.SubElement(svg, 'text')
        tx.set('x', str(x + w // 2))
        tx.set('y', str(y + h // 2 + 5))
        tx.set('text-anchor', 'middle')
        tx.set('font-size', str(FONT_SIZE))
        tx.set('font-family', '標楷體')
        tx.text = text

    def arrow(x1, y1, x2, y2):
        ln = etree.SubElement(svg, 'line')
        ln.set('x1', str(x1))
        ln.set('y1', str(y1))
        ln.set('x2', str(x2))
        ln.set('y2', str(y2))
        ln.set('stroke', '#000')
        ln.set('stroke-width', '1.5')
        ln.set('marker-end', 'url(#arrow)')

    # Flow: 廠商提送 → 監造審查 → 符合? → 是:核定 / 否:退回
    cx = PAD_LEFT + BOX_W
    box(cx - BOX_W // 2, PAD_TOP, '廠商提送\n品質計畫')
    arrow(cx, PAD_TOP + BOX_H, cx, PAD_TOP + BOX_H + GAP_Y)
    box(cx - BOX_W // 2, PAD_TOP + BOX_H + GAP_Y, '監造單位\n審查')
    arrow(cx, PAD_TOP + 2 * (BOX_H + GAP_Y), cx, PAD_TOP + 2 * (BOX_H + GAP_Y) + GAP_Y)
    diamond(cx - DIA_W // 2, PAD_TOP + 2 * (BOX_H + GAP_Y) + GAP_Y, '符合？')

    # 是 →
    arrow(cx + DIA_W // 2, PAD_TOP + 2 * (BOX_H + GAP_Y) + GAP_Y + DIA_H // 2,
          cx + DIA_W // 2 + 80, PAD_TOP + 2 * (BOX_H + GAP_Y) + GAP_Y + DIA_H // 2)
    tx_yes = etree.SubElement(svg, 'text')
    tx_yes.set('x', str(cx + DIA_W // 2 + 40))
    tx_yes.set('y', str(PAD_TOP + 2 * (BOX_H + GAP_Y) + GAP_Y + DIA_H // 2 - 8))
    tx_yes.set('text-anchor', 'middle')
    tx_yes.set('font-size', '11')
    tx_yes.text = '是'
    box(cx + DIA_W // 2 + 80, PAD_TOP + 2 * (BOX_H + GAP_Y) + GAP_Y + DIA_H // 2 - BOX_H // 2,
        '核定', BOX_W // 2, BOX_H)

    # 否 →
    arrow(cx, PAD_TOP + 2 * (BOX_H + GAP_Y) + GAP_Y + DIA_H,
          cx, PAD_TOP + 2 * (BOX_H + GAP_Y) + GAP_Y + DIA_H + GAP_Y)
    tx_no = etree.SubElement(svg, 'text')
    tx_no.set('x', str(cx + 15))
    tx_no.set('y', str(PAD_TOP + 2 * (BOX_H + GAP_Y) + GAP_Y + DIA_H + GAP_Y // 2))
    tx_no.set('text-anchor', 'middle')
    tx_no.set('font-size', '11')
    tx_no.text = '否'
    box(cx - BOX_W // 2, PAD_TOP + 2 * (BOX_H + GAP_Y) + GAP_Y + DIA_H + GAP_Y,
        '退回補正\n或重送')

    tree = etree.ElementTree(svg)
    tree.write(svg_path, encoding='utf-8', xml_declaration=True)
    print(f'  SVG 流程圖 → {os.path.basename(svg_path)}')


def add_review_table(doc):
    """加入表3.1 審查意見表（簡化版）"""
    doc.add_heading('表3.1  品質計畫審查意見表', level=3)
    items = [
        ('一、計畫範圍', '工程概要、主要施工項目、適用對象、名詞定義'),
        ('二、管理權責及分工', '品管組織架構、人員配置及職掌'),
        ('三、施工要領', '分項工程施工要領項目、內容大綱'),
        ('四、品質管理標準', '分項工程品管標準項目、標準化表單'),
        ('五、材料與設備檢驗', '送審資料、試驗室資格、自主檢查程序'),
        ('六、自主檢查表', '分項自主檢查項目、標準化表單'),
        ('七、不合格品管制', '材料及施工不合格管制程序'),
        ('八、矯正與預防措施', '矯正及預防措施辦理時機與流程'),
        ('九、內部品質稽核', '稽核範圍、頻率、缺失列管'),
        ('十、文件紀錄管理', '文件及紀錄管理作業程序、歸檔規劃'),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['項次', '審查項目', '審查重點']):
        hdr[i].text = txt
    for item, desc in items:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = desc.split('、')[0] if '、' in desc else desc
        row[2].text = desc


def build_doc(content_path, output_path, svg_output_path):
    sections = read_content_md(content_path)
    gen_flow_chart_svg(svg_output_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    doc.add_heading('第三章  品質計畫審查作業程序', level=1)

    # ——— 1. 審查作業程序 ———
    doc.add_heading('1  審查作業程序', level=2)

    doc.add_heading('(1) 審查核定流程', level=3)
    doc.add_paragraph('品質計畫之審查及核定流程如下圖所示。')
    doc.add_paragraph(f'（SVG 流程圖：{os.path.basename(svg_output_path)}）')

    doc.add_heading('(2) 審查時限', level=3)
    doc.add_paragraph('廠商提送品質計畫後，監造單位應於契約規定期限內完成審查。')

    doc.add_heading('(3) 不符合處理', level=3)
    for line in ['輕微不符：退回補正，限期 7 日內補送。',
                 '重大不符：退回重送，限期 14 日內重送。',
                 '逾期未補送或重送：依契約規定辦理。']:
        doc.add_paragraph(line, style='List Bullet')

    doc.add_heading('(4) 送審管制', level=3)
    doc.add_paragraph('品質計畫送審情形應列表管制，掌握各版本送審、退回、核定之狀態。')

    doc.add_heading('(5) 品管人員審查核定', level=3)
    doc.add_paragraph('對廠商品管組織人員之審查及核定作業，包含品管人員資格審查、'
                      '更換及補員期限等相關作業規定。')

    # ——— 2. 審查重點 ———
    doc.add_heading('2  審查重點', level=2)
    doc.add_paragraph('依契約及相關規定，訂定品質計畫審查重點如下表。')
    add_review_table(doc)

    # ——— 3. 應用表單 ———
    doc.add_heading('3  應用表單', level=2)
    doc.add_paragraph('本章使用之表單如下：', style='List Bullet')
    doc.add_paragraph('表3.1  品質計畫審查意見表', style='List Bullet')

    doc.save(output_path)
    print(f'Ch3 品質計畫審查作業程序 → {os.path.basename(output_path)}')


def main():
    parser = argparse.ArgumentParser(description='Ch3 品質計畫審查作業程序轉換工具 v1.0')
    parser.add_argument('-c', '--content', default='./ch3_內容.md')
    parser.add_argument('-o', '--output', default='../../output/Ch3_品質計畫審查作業程序.docx')
    parser.add_argument('--svg-output', default='../../output/Ch3_審查流程圖.svg')
    args = parser.parse_args()
    base = os.path.dirname(os.path.abspath(__file__))
    build_doc(
        os.path.normpath(os.path.join(base, args.content)),
        os.path.normpath(os.path.join(base, args.output)),
        os.path.normpath(os.path.join(base, args.svg_output)))


if __name__ == '__main__':
    main()
