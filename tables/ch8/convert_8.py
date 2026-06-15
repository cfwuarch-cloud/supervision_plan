# -*- coding: utf-8 -*-
"""
品質稽核 (Ch8) 轉換工具 v1.0
==============================
產出監造計畫第八章「品質稽核」docx，含流程圖 SVG 及稽核查對表。

修正歷程：
  v1.0  2026/06/15  初始版本

作者：OpenCode Assistant / cfwuarch
版本：v1.0
最後更新：2026/06/15

使用方法：
  python -X utf8 tables/ch8/convert_8.py

參數說明：
  -c, --content    文字內容 md 路徑
  -o, --output     輸出 docx 路徑
  --svg-output     SVG 流程圖輸出路徑
"""
import argparse
import os
import sys
import re
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))))


def read_content_md(path, placeholders=None):
    sections = {}
    if not path or not os.path.isfile(path):
        return sections
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    if placeholders:
        for k, v in placeholders.items():
            text = text.replace('{{' + k + '}}', v)
    current_heading = None
    current_lines = []
    for line in text.split('\n'):
        if line.startswith('# ') or line.startswith('## '):
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_lines).strip()
            current_heading = re.sub(r'^#+ ', '', line).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_lines).strip()
    return sections


def gen_flow_chart_svg(svg_path):
    """生成品質稽核流程 SVG"""
    svg_w, svg_h = 600, 200
    nsmap = {None: 'http://www.w3.org/2000/svg'}
    svg = etree.Element('svg', nsmap=nsmap)
    svg.set('width', str(svg_w))
    svg.set('height', str(svg_h))
    svg.set('viewBox', f'0 0 {svg_w} {svg_h}')

    defs = etree.SubElement(svg, 'defs')
    marker = etree.SubElement(defs, 'marker')
    marker.set('id', 'arrow')
    marker.set('markerWidth', '10')
    marker.set('markerHeight', '10')
    marker.set('refX', '10')
    marker.set('refY', '5')
    marker.set('orient', 'auto')
    pm = etree.SubElement(marker, 'path')
    pm.set('d', 'M0,0 L10,5 L0,10 Z')
    pm.set('fill', '#000')

    steps = ['通知', '起始會議', '現場稽核', '後會議', '矯正預防', '結案']
    BOX_W, BOX_H = 70, 40
    GAP = 15
    PAD_LEFT, PAD_TOP = 20, 60

    for i, text in enumerate(steps):
        x = PAD_LEFT + i * (BOX_W + GAP)
        y = PAD_TOP
        r = etree.SubElement(svg, 'rect')
        r.set('x', str(x))
        r.set('y', str(y))
        r.set('width', str(BOX_W))
        r.set('height', str(BOX_H))
        r.set('fill', '#fff')
        r.set('stroke', '#000')
        r.set('stroke-width', '1.5')
        tx = etree.SubElement(svg, 'text')
        tx.set('x', str(x + BOX_W // 2))
        tx.set('y', str(y + BOX_H // 2 + 5))
        tx.set('text-anchor', 'middle')
        tx.set('font-size', '12')
        tx.set('font-family', '標楷體')
        tx.text = text

        if i < len(steps) - 1:
            nx = PAD_LEFT + (i + 1) * (BOX_W + GAP)
            ln = etree.SubElement(svg, 'line')
            ln.set('x1', str(x + BOX_W))
            ln.set('y1', str(y + BOX_H // 2))
            ln.set('x2', str(nx))
            ln.set('y2', str(y + BOX_H // 2))
            ln.set('stroke', '#000')
            ln.set('stroke-width', '1.5')
            ln.set('marker-end', 'url(#arrow)')

    tree = etree.ElementTree(svg)
    tree.write(svg_path, encoding='utf-8', xml_declaration=True)
    print(f'  SVG 流程圖 → {os.path.basename(svg_path)}')


def build_doc(content_path, output_path, svg_output_path):
    sections = read_content_md(content_path)
    gen_flow_chart_svg(svg_output_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    doc.add_heading('第八章  品質稽核', level=1)

    # ——— 1. 品質稽核權責 ———
    doc.add_heading('1  品質稽核權責', level=2)
    doc.add_paragraph('監造人員執行品質稽核之權責說明如下：')
    for item in ['監造主任：主持內外部稽核作業。',
                 '品管人員：執行稽核、撰寫稽核報告。',
                 '監造工程師：配合受稽並提供相關資料。']:
        doc.add_paragraph(item, style='List Bullet')

    # ——— 2. 品質稽核範圍 ———
    doc.add_heading('2  品質稽核範圍', level=2)

    doc.add_heading('(1) 外部稽核（對廠商）', level=3)
    for item in ['廠商品質管理系統運作情形。',
                 '施工現場品質管制落實度。',
                 '自主檢查執行狀況。']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('(2) 內部稽核（監造自我）', level=3)
    for item in ['監造人員執行抽查抽驗之落實度。',
                 '文件紀錄管理之完整性。',
                 '缺失改善追蹤之閉合情形。']:
        doc.add_paragraph(item, style='List Bullet')

    # ——— 3. 品質稽核頻率 ———
    doc.add_heading('3  品質稽核頻率', level=2)
    doc.add_paragraph('定期稽核：每季至少一次。')
    doc.add_paragraph('不定期稽核：依工程執行情形適時辦理。')

    # ——— 4. 品質稽核流程 ———
    doc.add_heading('4  品質稽核流程', level=2)
    doc.add_paragraph('稽核流程如下圖所示。')
    doc.add_paragraph(f'（SVG 流程圖：{os.path.basename(svg_output_path)}）')
    doc.add_paragraph('')
    doc.add_paragraph('稽核流程包含：稽核之通知 → 起始會議 → 現場稽核 → '
                      '稽核後會議 → 稽核結果通知 → 矯正及預防措施 → 結案。')

    # ——— 稽核查對表 ———
    doc.add_heading('品質稽核查對表（參考例）', level=3)
    items = [
        ('組織與人員', '品管組織是否健全、人員配置是否合約規定'),
        ('文件審查', '品質計畫、施工計畫是否已審查核定'),
        ('材料檢驗', '材料送審是否完成、檢驗報告是否齊全'),
        ('施工抽查', '抽查紀錄是否完整、檢驗停留點是否落實'),
        ('不合格管制', '不合格品是否列管追蹤、改善是否閉合'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '稽核項目'
    hdr[1].text = '稽核重點'
    for item, desc in items:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = desc

    # ——— 5. 應用表單 ———
    doc.add_heading('5  應用表單', level=2)
    for f in ['內部稽核查對表', '外部稽核查對表', '稽核時程計畫管制表']:
        doc.add_paragraph(f, style='List Bullet')

    doc.save(output_path)
    print(f'Ch8 品質稽核 → {os.path.basename(output_path)}')


def main():
    parser = argparse.ArgumentParser(description='Ch8 品質稽核轉換工具 v1.0')
    parser.add_argument('-c', '--content', default='./ch8_內容.md')
    parser.add_argument('-o', '--output', default='../../output/Ch8_品質稽核.docx')
    parser.add_argument('--svg-output', default='../../output/Ch8_稽核流程圖.svg')
    args = parser.parse_args()
    base = os.path.dirname(os.path.abspath(__file__))
    build_doc(
        os.path.normpath(os.path.join(base, args.content)),
        os.path.normpath(os.path.join(base, args.output)),
        os.path.normpath(os.path.join(base, args.svg_output)))


if __name__ == '__main__':
    main()
