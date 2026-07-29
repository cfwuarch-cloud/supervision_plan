# -*- coding: utf-8 -*-
"""
Word 表格內容檢視工具
=====================
將 docx 中指定頁（表格）的內容以管線分隔格式輸出到終端機，
列號在左、每格標示 Cx:內容，格線分隔易於定位。

使用方法：
  python show_table.py -i 輸入.docx
  python show_table.py -i 輸入.docx -p 3
  python show_table.py -i 輸入.docx -p 1 -w 12 -r 0-5

參數：
  -i, --input      輸入 docx 路徑（必要）
  -p, --page       表格編號，從 1 開始（預設：1）
  -w, --width      每欄最大字數（預設：15，-1=不截斷）
  -r, --row-range  列範圍，例如 0-5 或 2- （預設：全部）
"""
import argparse
import os
import sys
import docx
from docx.oxml.ns import qn


def get_cell_text(tc, width):
    t = tc.text if tc.text else ''
    t = t.replace('\n', ' ').replace('\r', ' ').strip()
    if width > 0 and len(t) > width:
        t = t[:width - 1] + '~'
    return t


def get_vmerge(tc):
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        return ''
    vm = tcPr.find(qn('w:vMerge'))
    if vm is None:
        return ''
    val = vm.get(qn('w:val'))
    return 'M' if val == 'restart' else 'm'


def parse_row_range(row_range_str, max_rows):
    row_start, row_end = 0, max_rows
    if not row_range_str:
        return row_start, row_end
    try:
        parts = row_range_str.split('-')
        if parts[0]:
            row_start = max(0, int(parts[0]))
        if len(parts) > 1 and parts[1]:
            row_end = min(max_rows, int(parts[1]) + 1)
    except ValueError:
        print(f'警告：列範圍格式錯誤 "{row_range_str}"，顯示全部列。')
        return 0, max_rows
    return row_start, row_end


def show_table():
    parser = argparse.ArgumentParser(description='Word 表格內容檢視工具')
    parser.add_argument('-i', '--input', required=True, help='輸入 docx 路徑')
    parser.add_argument('-p', '--page', type=int, default=1, help='表格編號（1起，預設：1）')
    parser.add_argument('-w', '--width', type=int, default=15, help='每欄最大字數（預設：15，-1=不截斷）')
    parser.add_argument('-r', '--row-range', help='列範圍，例如 0-5 或 2- （預設：全部）')
    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f'錯誤：檔案不存在 {args.input}')
        sys.exit(1)

    try:
        doc = docx.Document(args.input)
    except Exception as e:
        print(f'錯誤：無法讀取檔案 ({e})')
        sys.exit(1)

    if not doc.tables:
        print('錯誤：該檔案中沒有找到任何表格。')
        sys.exit(1)

    if args.page < 1 or args.page > len(doc.tables):
        print(f'錯誤：表格編號 {args.page} 超出範圍（1~{len(doc.tables)}）')
        sys.exit(1)

    tbl = doc.tables[args.page - 1]
    trs = tbl._tbl.findall(qn('w:tr'))

    if not trs:
        print(f'表格 {args.page} 是空的。')
        return

    row_start, row_end = parse_row_range(args.row_range, len(trs))

    max_cols = 0
    for tr in trs[row_start:row_end]:
        n = len(tr.findall(qn('w:tc')))
        if n > max_cols:
            max_cols = n

    print()
    print(f'表格 {args.page}  R{row_start}~R{row_end-1}  |  {row_end - row_start}列 x {max_cols}欄')
    print('M=合併起始  m=合併延續')
    print('=' * 60)

    for ri in range(row_start, row_end):
        tr = trs[ri]
        tcs = tr.findall(qn('w:tc'))
        parts = [f'R{ri:02d}']
        for ci in range(max_cols):
            if ci < len(tcs):
                txt = get_cell_text(tcs[ci], args.width)
                vm = get_vmerge(tcs[ci])
                tag = f'C{ci}{vm}' if vm else f'C{ci}'
                parts.append(f'{tag}:{txt}')
            else:
                parts.append(f'C{ci}:')
        line = ' | '.join(parts)
        print(line)
        print('-' * min(len(line), 100))

    print()


if __name__ == '__main__':
    show_table()