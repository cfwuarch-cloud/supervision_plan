# -*- coding: utf-8 -*-
"""
施工計畫審查作業程序 (Ch4) 轉換工具 v1.0
==========================================
產出監造計畫第四章 docx，含審查流程圖 SVG 及表4.1/4.2 審查重點表。

修正歷程：
  v1.0  2026/06/15  初始版本

作者：OpenCode Assistant / cfwuarch
版本：v1.0
最後更新：2026/06/15

相依套件：
  - python-docx>=1.0.0,<2.0.0
  - lxml>=4.9.0,<6.0.0

使用方法：
  python -X utf8 tables/ch4/convert_4.py

參數說明：
  -c, --content    文字內容 md 路徑
  -o, --output     輸出 docx 路徑
  --svg-output     SVG 流程圖輸出路徑
"""
import argparse
import os
import sys
import re
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))))


def read_content_md(path, placeholders=None):
    sections = {}
    if not path or not os.path.isfile(path):
        return sections
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    if placeholders:
        for k, v in placeholders.items():
            text = text.replace('{{' + k + '}}', v)
    current_heading = None
    current_lines = []
    for line in text.split('\n'):
        if line.startswith('# ') or line.startswith('## '):
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_lines).strip()
            current_heading = re.sub(r'^#+ ', '', line).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_lines).strip()
    return sections


def gen_flow_chart_svg(svg_path):
    """生成施工計畫審查流程 SVG（14px 底部間距規範）"""
    svg_w, svg_h = 400, 500
    nsmap = {None: 'http://www.w3.org/2000/svg'}
    svg = etree.Element('svg', nsmap=nsmap)
    svg.set('width', str(svg_w))
    svg.set('height', str(svg_h))
    svg.set('viewBox', f'0 0 {svg_w} {svg_h}')

    defs = etree.SubElement(svg, 'defs')
    marker = etree.SubElement(defs, 'marker')
    marker.set('id', 'arrow')
    marker.set('markerWidth', '10')
    marker.set('markerHeight', '10')
    marker.set('refX', '10')
    marker.set('refY', '5')
    marker.set('orient', 'auto')
    pm = etree.SubElement(marker, 'path')
    pm.set('d', 'M0,0 L10,5 L0,10 Z')
    pm.set('fill', '#000')

    PAD_LEFT, PAD_TOP = 30, 30
    BOX_W, BOX_H = 180, 45
    GAP_Y = 25

    steps = [
        '廠商提送\n施工計畫',
        '監造單位\n審查',
        '符合？\n（核定）',
        '退回補正\n或重送',
    ]

    y_positions = [PAD_TOP + i * (BOX_H + GAP_Y) for i in range(4)]
    cx = PAD_LEFT + BOX_W // 2 + 40

    for i, text in enumerate(steps):
        y = y_positions[i]
        if i == 2:
            # 菱形
            d = etree.SubElement(svg, 'polygon')
            pts = f'{cx},{y} {cx+BOX_W//2},{y+BOX_H//2} {cx},{y+BOX_H} {cx-BOX_W//2},{y+BOX_H//2}'
            d.set('points', pts)
            d.set('fill', '#fff')
            d.set('stroke', '#000')
            d.set('stroke-width', '1.5')
            tx = etree.SubElement(svg, 'text')
            tx.set('x', str(cx))
            tx.set('y', str(y + BOX_H // 2 + 5))
            tx.set('text-anchor', 'middle')
            tx.set('font-size', '13')
            tx.set('font-family', '標楷體')
            tx.text = '符合？'
        else:
            r = etree.SubElement(svg, 'rect')
            r.set('x', str(cx - BOX_W // 2))
            r.set('y', str(y))
            r.set('width', str(BOX_W))
            r.set('height', str(BOX_H))
            r.set('fill', '#fff')
            r.set('stroke', '#000')
            r.set('stroke-width', '1.5')
            lines = text.split('\n')
            for li, t in enumerate(lines):
                tx = etree.SubElement(svg, 'text')
                tx.set('x', str(cx))
                tx.set('y', str(y + BOX_H // 2 - (len(lines)-1)*9 + li*18 + 5))
                tx.set('text-anchor', 'middle')
                tx.set('font-size', '13')
                tx.set('font-family', '標楷體')
                tx.text = t

        if i < len(steps) - 1:
            ny = y_positions[i + 1]
            if i == 0:
                ln = etree.SubElement(svg, 'line')
                ln.set('x1', str(cx))
                ln.set('y1', str(y + BOX_H))
                ln.set('x2', str(cx))
                ln.set('y2', str(ny))
                ln.set('stroke', '#000')
                ln.set('stroke-width', '1.5')
                ln.set('marker-end', 'url(#arrow)')
            elif i == 1:
                ln = etree.SubElement(svg, 'line')
                ln.set('x1', str(cx))
                ln.set('y1', str(y + BOX_H))
                ln.set('x2', str(cx))
                ln.set('y2', str(ny))
                ln.set('stroke', '#000')
                ln.set('stroke-width', '1.5')
                ln.set('marker-end', 'url(#arrow)')
            elif i == 2:
                # 向下（否）
                ln = etree.SubElement(svg, 'line')
                ln.set('x1', str(cx))
                ln.set('y1', str(y + BOX_H))
                ln.set('x2', str(cx))
                ln.set('y2', str(ny))
                ln.set('stroke', '#000')
                ln.set('stroke-width', '1.5')
                ln.set('marker-end', 'url(#arrow)')
                tx_no = etree.SubElement(svg, 'text')
                tx_no.set('x', str(cx + 15))
                tx_no.set('y', str(y + BOX_H + GAP_Y // 2))
                tx_no.set('text-anchor', 'middle')
                tx_no.set('font-size', '11')
                tx_no.text = '否'

                # 向右（是）
                ln2 = etree.SubElement(svg, 'line')
                ln2.set('x1', str(cx))
                ln2.set('y1', str(y + BOX_H // 2))
                ln2.set('x2', str(cx + BOX_W))
                ln2.set('y2', str(y + BOX_H // 2))
                ln2.set('stroke', '#000')
                ln2.set('stroke-width', '1.5')
                ln2.set('marker-end', 'url(#arrow)')
                tx_yes = etree.SubElement(svg, 'text')
                tx_yes.set('x', str(cx + BOX_W // 2))
                tx_yes.set('y', str(y + BOX_H // 2 - 8))
                tx_yes.set('text-anchor', 'middle')
                tx_yes.set('font-size', '11')
                tx_yes.text = '是'

                # 核定框
                r2 = etree.SubElement(svg, 'rect')
                r2.set('x', str(cx + BOX_W))
                r2.set('y', str(y + BOX_H // 2 - BOX_H // 2))
                r2.set('width', str(BOX_W // 2))
                r2.set('height', str(BOX_H))
                r2.set('fill', '#fff')
                r2.set('stroke', '#000')
                r2.set('stroke-width', '1.5')
                tx2 = etree.SubElement(svg, 'text')
                tx2.set('x', str(cx + BOX_W + BOX_W // 4))
                tx2.set('y', str(y + BOX_H // 2 + 5))
                tx2.set('text-anchor', 'middle')
                tx2.set('font-size', '13')
                tx2.set('font-family', '標楷體')
                tx2.text = '核定'

    tree = etree.ElementTree(svg)
    tree.write(svg_path, encoding='utf-8', xml_declaration=True)
    print(f'  SVG 流程圖 → {os.path.basename(svg_path)}')


def add_table_41(doc):
    """整體施工計畫審查重點表（表4.1）"""
    doc.add_heading('表4.1  整體施工計畫審查重點表', level=3)
    items = [
        ('計畫書架構', '計畫書內容與工程契約相關規定是否相符'),
        ('一、工程概要', '主要施工項目、材料規格或工法、相關數量'),
        ('二、開工前置作業', '地質資料研判、工址調查、障礙物處理'),
        ('三、施工作業管理', '工地組織、施工機具、施工程序規劃'),
        ('四、進度管理', '施工預定進度、要徑作業、協調會議'),
        ('五、假設工程計畫', '工區配置、整地計畫、臨時設施規劃'),
        ('六、施工測量', '控制測量、施工測量方法'),
        ('七、施工區域排水', '排水系統調查、擋水抽水措施'),
        ('八、分項施工計畫', '分項計畫項目、提送時程'),
        ('九、職業安全衛生', '安衛組織、協議組織、教育訓練'),
        ('十、緊急應變及防災', '應變編組、通報系統、防災對策'),
        ('十一、環境保護', '環保組織、噪音防制、污染防治'),
        ('十二、交通維持及安全', '交維措施、運輸路線限制檢討'),
        ('十三、移交管理計畫', '移交文件項目、管理維護訓練計畫'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '項次'
    hdr[1].text = '審查項目'
    for item, desc in items:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = desc


def add_table_42(doc):
    """分項工程施工計畫審查重點表（表4.2）"""
    doc.add_heading('表4.2  分項工程施工計畫審查重點表', level=3)
    items = [
        ('一、工項概要', '分項工程概要說明、重要施作項目與數量'),
        ('二、人員組織', '必要人員配置、責任職掌、施工人數'),
        ('三、預定作業進度', '配合整體進度表、分項施工時程'),
        ('四、分項品質計畫', '施工要領、品質管理標準、檢驗程序、自主檢查表'),
        ('五、作業安全衛生', '勞安設施、人員管理'),
        ('六、施工圖說', '充分之施工圖或計算書'),
        ('七、相關附件', '協調會議紀錄、材料比對表、CNS規範'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '項次'
    hdr[1].text = '審查項目'
    for item, desc in items:
        row = table.add_row().cells
        row[0].text = item
        row[1].text = desc


def build_doc(content_path, output_path, svg_output_path):
    sections = read_content_md(content_path)
    gen_flow_chart_svg(svg_output_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    doc.add_heading('第四章  施工計畫審查作業程序', level=1)

    # ——— 1. 施工計畫分階段送審 ———
    doc.add_heading('1  施工計畫分階段送審', level=2)
    doc.add_paragraph('廠商應依契約規定，製作整體施工計畫及其他分項施工計畫，'
                      '並依整體施工預定進度表檢討訂定提送時限。')
    doc.add_paragraph('監造單位應明確條列廠商應送審之分項施工計畫項目，以利控管。')

    # ——— 2. 審查作業程序 ———
    doc.add_heading('2  審查作業程序', level=2)

    doc.add_heading('(1) 審查核定流程', level=3)
    doc.add_paragraph('施工計畫之審查及核定流程如下圖所示。')
    doc.add_paragraph(f'（SVG 流程圖：{os.path.basename(svg_output_path)}）')

    doc.add_heading('(2) 審查時限', level=3)
    doc.add_paragraph('施工計畫審查時限依契約規定辦理。')

    doc.add_heading('(3) 不符合處理', level=3)
    doc.add_paragraph('對於不符合情形，依輕微不符（補正）、重大不符（退回重送）分別處理。')

    doc.add_heading('(4) 送審管制', level=3)
    doc.add_paragraph('施工計畫送審過程應列表管制，重點包含送審及修改時程之掌控。')

    # ——— 3. 審查重點 ———
    doc.add_heading('3  審查重點', level=2)
    doc.add_paragraph('依契約內容，訂定整體施工計畫及分項施工計畫之審查重點如下。')
    add_table_41(doc)
    add_table_42(doc)

    # ——— 4. 應用表單 ———
    doc.add_heading('4  應用表單', level=2)
    for f in ['表4.1  整體施工計畫審查重點表', '表4.2  分項工程施工計畫審查重點表']:
        doc.add_paragraph(f, style='List Bullet')

    doc.save(output_path)
    print(f'Ch4 施工計畫審查作業程序 → {os.path.basename(output_path)}')


def main():
    parser = argparse.ArgumentParser(description='Ch4 施工計畫審查作業程序轉換工具 v1.0')
    parser.add_argument('-c', '--content', default='./ch4_內容.md')
    parser.add_argument('-o', '--output', default='../../output/Ch4_施工計畫審查作業程序.docx')
    parser.add_argument('--svg-output', default='../../output/Ch4_審查流程圖.svg')
    args = parser.parse_args()
    base = os.path.dirname(os.path.abspath(__file__))
    build_doc(
        os.path.normpath(os.path.join(base, args.content)),
        os.path.normpath(os.path.join(base, args.output)),
        os.path.normpath(os.path.join(base, args.svg_output)))


if __name__ == '__main__':
    main()
