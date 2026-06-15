# -*- coding: utf-8 -*-
"""
監造範圍 (Ch1) 轉換工具 v1.0
============================
將詳細價目表作為母本，讀取指定階層作為主要施工項目，
配合文字內容 markdown，產出監造計畫第一章 docx。

修正歷程：
  v1.0  2026/06/15  初始版本

作者：OpenCode Assistant / cfwuarch
版本：v1.0
最後更新：2026/06/15

相依套件：
  - openpyxl>=3.0.0,<4.0.0
  - python-docx>=1.0.0,<2.0.0
  - pandas>=1.5.0,<3.0.0

使用方法：
  python -X utf8 tables/ch1/convert_1.py --level 4

功能說明：
  1. 讀取詳細價目表（excel），依 level 參數控制階層深度
  2. 讀取 ch1_內容.md 文字內容，替換 {{placeholder}}
  3. 產出監造範圍章節 docx

參數說明：
  -c, --content    文字內容 md 路徑（預設：./ch1_內容.md）
  -p, --price      詳細價目表 Excel 路徑
  -o, --output     輸出 docx 路徑
  --level          價目表階層深度（預設 5，越大越細）
  --project-name   專案名稱（覆寫預設值）
"""
import argparse
import os
import sys
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))))
from common.docx_utils import add_page_break

PROJECT_NAME = '臺南市政府社會局委託辦理北區成德公設民營托嬰中心室內裝修統包工程'


def clean_text(text):
    text = text.replace('\n', '').replace('\r', '')
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[\s]+([、，,。．])', r'\1', text)
    return text.strip()


def load_price_sheet(path, level=5, exclude_units=None):
    """載入價目表，依 level 過濾階層深度（dot 數）"""
    if exclude_units is None:
        exclude_units = {'式', '工'}
    df = pd.read_excel(path, sheet_name='Table 1')
    c0 = df.columns[0]
    c1 = df.columns[1]
    cu = df.columns[6]
    cq = df.columns[7]

    items = []
    started = False
    for _, row in df.iterrows():
        item = clean_text(str(row[c0]))
        name = clean_text(str(row[c1]))
        unit = clean_text(str(row[cu])) if pd.notna(row[cu]) else ''
        qty = row[cq]

        name = name.replace('⑤', '窗')
        unit = unit.replace('⑤', '窗')

        if '壹.三.1' in item:
            started = True
        if not started:
            continue
        if not item.startswith('壹.'):
            continue
        if not unit or unit == 'nan':
            continue
        if unit in exclude_units:
            continue
        if any(k in name for k in ['小計', '合計', '總價']):
            continue
        if pd.isna(qty):
            continue

        depth = len(item.split('.'))
        if depth > level:
            continue

        if isinstance(qty, float) and qty == int(qty):
            qty_str = str(int(qty))
        else:
            qty_str = str(qty)
        items.append((item, name, qty_str + unit, depth))

    return items


def read_content_md(path, placeholders=None):
    """讀取 markdown 內容，替換 placeholder，回傳 {heading: text_body}"""
    sections = {}
    if not path or not os.path.isfile(path):
        return sections
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()

    if placeholders:
        for k, v in placeholders.items():
            text = text.replace('{{' + k + '}}', v)

    current_heading = None
    current_lines = []
    raw_lines = text.split('\n')
    for line in raw_lines:
        if line.startswith('# ') or line.startswith('## '):
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_lines).strip()
            current_heading = re.sub(r'^#+ ', '', line).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_lines).strip()

    return sections


def set_table_style(table):
    """設定表格邊框：所有儲存格加上黑實線框"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        be = OxmlElement(f'w:{edge}')
        be.set(qn('w:val'), 'single')
        be.set(qn('w:sz'), '4')
        be.set(qn('w:space'), '0')
        be.set(qn('w:color'), '000000')
        tblBorders.append(be)
    tblPr.append(tblBorders)


def build_doc(content_path, price_path, output_path, level=5, project_name=PROJECT_NAME):
    """產出 Ch1 監造範圍 docx"""
    # 讀取文字內容
    placeholders = {
        '專案名稱': project_name,
        '主辦機關': '{{主辦機關}}',
        '設計單位': '{{設計單位}}',
        '監造單位': '{{監造單位}}',
        '工程地點': '{{工程地點}}',
        '工程期限': '{{工程期限}}',
        '工程預算': '{{工程預算}}',
    }
    sections = read_content_md(content_path, placeholders)

    # 讀取價目表
    price_items = []
    if price_path and os.path.isfile(price_path):
        price_items = load_price_sheet(price_path, level=level)

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # ——— 章節標題 ———
    doc.add_heading('第一章  監造範圍', level=1)
    doc.add_paragraph()

    # ——— 1. 依據 ———
    doc.add_heading('1  依據', level=2)
    for line in sections.get('依據', '').split('\n'):
        line = line.strip()
        if line:
            if line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

    # ——— 2. 工程概要 ———
    doc.add_heading('2  工程概要', level=2)
    for line in sections.get('工程概要', '').split('\n'):
        line = line.strip()
        if line:
            if line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)

    # ——— 3. 工程主要施工項目及數量 ———
    doc.add_heading('3  工程主要施工項目及數量', level=2)
    doc.add_paragraph('有關契約中主要項目，包括數量較多或施工時程較長、金額較大、'
                      f'或使用特殊材料、規格、工法等，予以表列（階層深度 level={level}）。')

    if price_items:
        table = doc.add_table(rows=1, cols=3)
        set_table_style(table)

        hdr = table.rows[0].cells
        hdr[0].text = '項次'
        hdr[1].text = '項目名稱'
        hdr[2].text = '數量'
        for cell in hdr:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True

        for item, name, qty_unit, depth in price_items:
            row = table.add_row().cells
            row[0].text = item
            row[1].text = name
            row[2].text = qty_unit
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph('（價目表未載入或無符合項目）')

    # ——— 4. 適用對象 ———
    doc.add_heading('4  適用對象', level=2)
    for line in sections.get('適用對象', '').split('\n'):
        line = line.strip()
        if line:
            doc.add_paragraph(line)

    # ——— 5. 名詞定義 ———
    doc.add_heading('5  名詞定義', level=2)
    for line in sections.get('名詞定義', '').split('\n'):
        line = line.strip()
        if line:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f'Ch1 監造範圍 → {os.path.basename(output_path)}')
    print(f'  價目表項目：{len(price_items)} 項（level={level}）')


def main():
    parser = argparse.ArgumentParser(description='Ch1 監造範圍轉換工具 v1.0')
    parser.add_argument('-c', '--content', default='./ch1_內容.md',
                        help='文字內容 md 路徑')
    parser.add_argument('-p', '--price', default='../../data/02_成德-詳細價目表.xlsx',
                        help='詳細價目表 Excel 路徑')
    parser.add_argument('-o', '--output', default='../../output/Ch1_監造範圍.docx',
                        help='輸出 docx 路徑')
    parser.add_argument('--level', type=int, default=5,
                        help='價目表階層深度（預設 5）')
    parser.add_argument('--project-name', default=None,
                        help='專案名稱（覆寫預設值）')
    args = parser.parse_args()

    base = os.path.dirname(os.path.abspath(__file__))
    content_path = os.path.normpath(os.path.join(base, args.content))
    price_path = os.path.normpath(os.path.join(base, args.price))
    output_path = os.path.normpath(os.path.join(base, args.output))

    project_name = args.project_name or PROJECT_NAME

    build_doc(content_path, price_path, output_path,
              level=args.level, project_name=project_name)


if __name__ == '__main__':
    main()
