# -*- coding: utf-8 -*-
"""
文件紀錄管理系統 (Ch9) 轉換工具 v1.0
======================================
產出監造計畫第九章「文件紀錄管理系統」docx。

修正歷程：
  v1.0  2026/06/15  初始版本

作者：OpenCode Assistant / cfwuarch
版本：v1.0
最後更新：2026/06/15

使用方法：
  python -X utf8 tables/ch9/convert_9.py

參數說明：
  -c, --content    文字內容 md 路徑
  -o, --output     輸出 docx 路徑
"""
import argparse
import os
import sys
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))))


def read_content_md(path, placeholders=None):
    sections = {}
    if not path or not os.path.isfile(path):
        return sections
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    if placeholders:
        for k, v in placeholders.items():
            text = text.replace('{{' + k + '}}', v)
    current_heading = None
    current_lines = []
    for line in text.split('\n'):
        if line.startswith('# ') or line.startswith('## '):
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_lines).strip()
            current_heading = re.sub(r'^#+ ', '', line).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_lines).strip()
    return sections


def build_doc(content_path, output_path):
    sections = read_content_md(content_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    doc.add_heading('第九章  文件紀錄管理系統', level=1)

    # ——— 1. 文件管理系統 ———
    doc.add_heading('1  文件管理系統', level=2)
    doc.add_paragraph('對於與本工程所有相關文件項目詳予表列，並作適當之分類、編碼，'
                      '規劃其登錄、收發、核定、保存、作廢等作業程序及存放管理方式。')

    doc.add_heading('文件分類編碼', level=3)
    doc.add_paragraph('文件分類編碼原則如下表所示。')

    codes = [
        ('A', '契約文件', '工程契約、補充說明、變更設計'),
        ('B', '計畫書類', '品質計畫、施工計畫、監造計畫'),
        ('C', '送審資料', '材料送審、設備送審、施工圖說'),
        ('D', '檢試驗報告', '材料試驗報告、設備測試報告'),
        ('E', '查驗紀錄', '施工抽查紀錄、材料抽驗紀錄'),
        ('F', '會議紀錄', '協調會議、稽核會議'),
        ('G', '公文書信', '往來公文、函件、備忘錄'),
        ('H', '日報月報', '施工日報、監造日報、月報'),
        ('I', '照片紀錄', '施工照片、稽核查核照片'),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['分類代碼', '分類名稱', '說明']):
        hdr[i].text = txt
    for code, name, desc in codes:
        row = table.add_row().cells
        row[0].text = code
        row[1].text = name
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('文件編碼格式：分類代碼-序號-版次（例如：C-001-01）')

    # ——— 2. 紀錄管理作業程序 ———
    doc.add_heading('2  紀錄管理作業程序', level=2)
    doc.add_paragraph('規劃工地內所作各項相關紀錄資料之登錄、收發、核定、保存、作廢等'
                      '作業程序，及如何配合文件之分類、編碼等，'
                      '將其紀錄成果作有系統之歸檔。')
    doc.add_paragraph('紀錄管理應包含下列項目：')
    for item in ['各項查驗紀錄', '會議紀錄', '監造日報表', '施工日報表',
                 '自主檢查表', '內部稽核紀錄', '施工照片']:
        doc.add_paragraph(item, style='List Bullet')

    # ——— 3. 文件紀錄移轉及存檔 ———
    doc.add_heading('3  文件紀錄移轉及存檔', level=2)

    doc.add_heading('(1) 工程完工移轉', level=3)
    doc.add_paragraph('工程完工後，應將文件紀錄資料移轉予業主，包含：')
    for item in ['施工及監造相關紀錄', '檢試驗報告', '竣工圖說', '操作維護手冊']:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('(2) 存檔管理', level=3)
    doc.add_paragraph('文件紀錄資料最終之存檔位置及存檔年限依契約規定辦理。')

    doc.save(output_path)
    print(f'Ch9 文件紀錄管理系統 → {os.path.basename(output_path)}')


def main():
    parser = argparse.ArgumentParser(description='Ch9 文件紀錄管理系統轉換工具 v1.0')
    parser.add_argument('-c', '--content', default='./ch9_內容.md')
    parser.add_argument('-o', '--output', default='../../output/Ch9_文件紀錄管理系統.docx')
    args = parser.parse_args()
    base = os.path.dirname(os.path.abspath(__file__))
    build_doc(
        os.path.normpath(os.path.join(base, args.content)),
        os.path.normpath(os.path.join(base, args.output)))


if __name__ == '__main__':
    main()
