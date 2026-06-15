# -*- coding: utf-8 -*-
"""
設備功能運轉測試抽驗程序及標準 (Ch6) 轉換工具 v1.0
====================================================
產出監造計畫第六章（機電條件章節）docx，含表6.1、6.2。

修正歷程：
  v1.0  2026/06/15  初始版本

作者：OpenCode Assistant / cfwuarch
版本：v1.0
最後更新：2026/06/15

使用方法：
  python -X utf8 tables/ch6/convert_6.py

參數說明：
  -c, --content    文字內容 md 路徑
  -o, --output     輸出 docx 路徑
"""
import argparse
import os
import sys
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))))


def read_content_md(path, placeholders=None):
    sections = {}
    if not path or not os.path.isfile(path):
        return sections
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    if placeholders:
        for k, v in placeholders.items():
            text = text.replace('{{' + k + '}}', v)
    current_heading = None
    current_lines = []
    for line in text.split('\n'):
        if line.startswith('# ') or line.startswith('## '):
            if current_heading is not None:
                sections[current_heading] = '\n'.join(current_lines).strip()
            current_heading = re.sub(r'^#+ ', '', line).strip()
            current_lines = []
        else:
            current_lines.append(line)
    if current_heading is not None:
        sections[current_heading] = '\n'.join(current_lines).strip()
    return sections


def build_doc(content_path, output_path):
    sections = read_content_md(content_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    doc.add_heading('第六章  設備功能運轉測試抽驗程序及標準', level=1)
    doc.add_paragraph('（工程內含運轉類機電設備者應撰寫本章）')

    # ——— 1. 設備功能運轉測試抽驗程序 ———
    doc.add_heading('1  設備功能運轉測試抽驗程序', level=2)

    doc.add_heading('(1) 單機設備測試抽驗', level=3)
    doc.add_paragraph('為確認單機設備裝置能符合契約要求，依設備之性質，'
                      '檢討訂定抽驗作業程序及抽驗項目。')

    doc.add_heading('(2) 系統運轉測試抽驗', level=3)
    doc.add_paragraph('為確認機電整套系統設備相關之管路、電氣、儀控、監測等'
                      '裝配完成後之運作，能符合契約之要求，'
                      '依設備之性質，檢討訂定系統運轉抽驗項目。')

    doc.add_heading('(3) 整體功能試運轉抽驗', level=3)
    doc.add_paragraph('為確認所有機電設備系統相互連結後，整體之運作能符合契約之要求，'
                      '依設備之性質，檢討訂定整體功能試運轉抽驗項目及'
                      '承攬廠商應提交之記錄及報告。')

    # ——— 2. 設備功能運轉測試抽驗標準 ———
    doc.add_heading('2  設備功能運轉測試抽驗標準', level=2)
    doc.add_paragraph('對於各項設備功能運轉之檢驗，依單機、系統及整體功能運轉測試，'
                      '分別檢討訂定相關測試抽驗標準。')

    # ——— 表6.2 設備功能運轉抽驗標準表 ———
    doc.add_heading('表6.2  設備功能運轉抽驗標準表', level=3)
    standards = [
        ('單機測試', '型號規格、電壓、電流、馬力', '設備安裝完成後', '儀表量測、功能確認',
         '每台', '更換或修復', '測試紀錄表'),
        ('系統測試', '系統獨立功能、組合測試、排放測試', '系統組設完成後', '功能測試、管路檢查',
         '每次', '檢修後複測', '系統測試報告'),
        ('整體測試', '整體系統連結整合功能', '全系統完成後', '試運轉、功能驗證',
         '至少一次', '改善後重新測試', '試運轉報告'),
    ]
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(['測試流程', '管理項目', '管理標準', '抽驗時機', '抽驗方法',
                              '抽驗頻率', '不合格處理']):
        hdr[i].text = txt
    for row_data in standards:
        row = table.add_row().cells
        for i, txt in enumerate(row_data):
            row[i].text = txt

    # ——— 表6.1 設備功能運轉測試紀錄表 ———
    doc.add_heading('表6.1  設備功能運轉測試紀錄表', level=3)
    table2 = doc.add_table(rows=4, cols=4)
    table2.style = 'Table Grid'
    table2.rows[0].cells[0].text = '工程名稱'
    table2.rows[0].cells[1].merge(table2.rows[0].cells[3])
    table2.rows[0].cells[1].text = '（工程名稱）'
    table2.rows[1].cells[0].text = '分項工程名稱'
    table2.rows[1].cells[1].merge(table2.rows[1].cells[3])
    table2.rows[1].cells[1].text = '（分項工程名稱）'
    table2.rows[2].cells[0].text = '測試流程'
    table2.rows[2].cells[1].merge(table2.rows[2].cells[3])
    table2.rows[2].cells[1].text = '□ 單機測試  □ 系統測試  □ 整體功能運轉測試'
    table2.rows[3].cells[0].text = '抽驗位置'
    table2.rows[3].cells[1].text = '（位置）'
    table2.rows[3].cells[2].text = '抽驗日期'
    table2.rows[3].cells[3].text = '（日期）'

    doc.add_paragraph('抽驗項目、抽驗標準及結果欄位請依實際設備特性自行填列。')

    # ——— 3. 應用表單 ———
    doc.add_heading('3  應用表單', level=2)
    for f in ['表6.1  設備功能運轉測試紀錄表', '表6.2  設備功能運轉抽驗標準表']:
        doc.add_paragraph(f, style='List Bullet')

    doc.save(output_path)
    print(f'Ch6 設備功能運轉測試抽驗程序及標準 → {os.path.basename(output_path)}')


def main():
    parser = argparse.ArgumentParser(
        description='Ch6 設備功能運轉測試抽驗程序及標準轉換工具 v1.0')
    parser.add_argument('-c', '--content', default='./ch6_內容.md')
    parser.add_argument('-o', '--output', default='../../output/Ch6_設備功能運轉測試抽驗程序及標準.docx')
    args = parser.parse_args()
    base = os.path.dirname(os.path.abspath(__file__))
    build_doc(
        os.path.normpath(os.path.join(base, args.content)),
        os.path.normpath(os.path.join(base, args.output)))


if __name__ == '__main__':
    main()
