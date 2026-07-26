# -*- coding: utf-8 -*-
"""表5-1 V2.6 — 採用 convert_5.1_v2.py 的表格格式與文字型式
"""
import sys, os
sys.path.insert(0, r'C:\OpenCodeLib\supervision_plan')
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from copy import deepcopy
from common.docx_table import add_cell

TEMPLATE = r'C:\OpenCodeLib\supervision_plan\tables\table5.1\表5.1.docx'
DATA_SRC = r'C:\OpenCodeLib\supervision_plan\data\(安南托育)表5-1_材料設備送審管制總表_0715.docx'
OUT = r'C:\OpenCodeLib\supervision_plan\output\(安南托育)表5-1_材料設備送審管制總表V2.6.docx'

PAGE_PLAN = [12, 11, 10, 10, 12, 11, 10, 10, 12, 13, 7]
DAT_MERGE = {0, 2, 3, 6, 7, 8, 9, 10, 11, 12, 14}
COL_W = [288, 2899, 857, 499, 1185, 547, 438, 488, 394, 455, 404, 434, 335, 865, 438]
LINE_H_TWIP = 240
PROJECT = '臺南市政府社會局委託辦理北區成德公設民營托嬰中心室內裝修統包工程'

# ===== Phase 1: 讀模板 & 資料 =====
tmpl = Document(TEMPLATE)
body = tmpl.element.body
t0 = tmpl.tables[0]
title_trs = [deepcopy(tr) for tr in t0._tbl.findall(qn('w:tr'))[:2]]
sect_pr = deepcopy(next(c for c in reversed(list(body)) if c.tag == qn('w:sectPr')))

src = Document(DATA_SRC)
src_body = src.element.body
src_tbls = [c for c in list(src_body) if c.tag == qn('w:tbl')]
raw_pairs = []
for ti, tbl_el in enumerate(src_tbls):
    trs = list(tbl_el.findall(qn('w:tr')))
    nh = 3 if ti == 0 else 2
    for i in range(nh, len(trs), 2):
        if i+1 >= len(trs): break
        raw_pairs.append((deepcopy(trs[i]), deepcopy(trs[i+1])))
assert len(raw_pairs) >= sum(PAGE_PLAN) + 4, f'{len(raw_pairs)} < {sum(PAGE_PLAN) + 4}'
# 排除 _0715 中多出的 4 項以匹配原始 V2
SKIP = {112, 113, 117, 118}
raw_pairs = [p for pi, p in enumerate(raw_pairs) if pi not in SKIP]
assert len(raw_pairs) == sum(PAGE_PLAN), f'{len(raw_pairs)} != {sum(PAGE_PLAN)}'
print(f'Loaded {len(raw_pairs)} pairs (from _0715, skipped 4)')

# ===== Phase 2: 清空 body =====
for c in list(body):
    if c.tag != qn('w:sectPr'): body.remove(c)

def copy_text(tr, ci):
    tcs = tr.findall(qn('w:tc'))
    if ci < len(tcs): return ''.join(t.text or '' for t in tcs[ci].iter(qn('w:t')))
    return ''

def add_header(body, pgn, ttl):
    """v2 風格表頭三行"""
    p1 = etree.SubElement(body, qn('w:p'))
    pPr1 = etree.SubElement(p1, qn('w:pPr'))
    jc1 = etree.SubElement(pPr1, qn('w:jc')); jc1.set(qn('w:val'),'center')
    sp1 = etree.SubElement(pPr1, qn('w:spacing')); sp1.set(qn('w:line'),'480'); sp1.set(qn('w:lineRule'),'exact')
    r1 = etree.SubElement(p1, qn('w:r'))
    rPr1 = etree.SubElement(r1, qn('w:rPr'))
    rf1 = etree.SubElement(rPr1, qn('w:rFonts')); rf1.set(qn('w:ascii'),'Arial'); rf1.set(qn('w:hAnsi'),'Arial'); rf1.set(qn('w:eastAsia'),'標楷體')
    etree.SubElement(rPr1, qn('w:b'))
    etree.SubElement(rPr1, qn('w:sz')).set(qn('w:val'),'28')
    etree.SubElement(rPr1, qn('w:szCs')).set(qn('w:val'),'28')
    t1 = etree.SubElement(r1, qn('w:t')); t1.text = f'表5-1 材料設備送審管制總表-{pgn}'

    p2 = etree.SubElement(body, qn('w:p'))
    pPr2 = etree.SubElement(p2, qn('w:pPr'))
    jc2 = etree.SubElement(pPr2, qn('w:jc')); jc2.set(qn('w:val'),'left')
    r2 = etree.SubElement(p2, qn('w:r'))
    rPr2 = etree.SubElement(r2, qn('w:rPr'))
    rf2 = etree.SubElement(rPr2, qn('w:rFonts')); rf2.set(qn('w:ascii'),'Arial'); rf2.set(qn('w:hAnsi'),'Arial'); rf2.set(qn('w:eastAsia'),'標楷體')
    etree.SubElement(rPr2, qn('w:sz')).set(qn('w:val'),'24')
    etree.SubElement(rPr2, qn('w:szCs')).set(qn('w:val'),'24')
    t2 = etree.SubElement(r2, qn('w:t')); t2.text = f'工程名稱：{PROJECT}'

    p3 = etree.SubElement(body, qn('w:p'))
    pPr3 = etree.SubElement(p3, qn('w:pPr'))
    jc3 = etree.SubElement(pPr3, qn('w:jc')); jc3.set(qn('w:val'),'left')
    sp3 = etree.SubElement(pPr3, qn('w:spacing')); sp3.set(qn('w:line'),'240'); sp3.set(qn('w:lineRule'),'atLeast')
    r3 = etree.SubElement(p3, qn('w:r'))
    rPr3 = etree.SubElement(r3, qn('w:rPr'))
    rf3 = etree.SubElement(rPr3, qn('w:rFonts')); rf3.set(qn('w:ascii'),'Times New Roman'); rf3.set(qn('w:hAnsi'),'Times New Roman'); rf3.set(qn('w:eastAsia'),'標楷體')
    etree.SubElement(rPr3, qn('w:sz')).set(qn('w:val'),'24')
    etree.SubElement(rPr3, qn('w:szCs')).set(qn('w:val'),'24')
    t3 = etree.SubElement(r3, qn('w:t')); t3.text = f'(監造單位使用)                              第{pgn}頁共{ttl}頁   表單編號：E51-{pgn}'

# ===== Phase 3: 組裝 =====
pair_idx = 0
for pg, n_pairs in enumerate(PAGE_PLAN):
    pgn = pg + 1; ttl = len(PAGE_PLAN)

    if pg > 0:
        pb = etree.SubElement(body, qn('w:p'))
        pr = etree.SubElement(pb, qn('w:r'))
        bb = etree.SubElement(pr, qn('w:br')); bb.set(qn('w:type'),'page')

    add_header(body, pgn, ttl)

    tbl = etree.SubElement(body, qn('w:tbl'))
    tblPr = etree.SubElement(tbl, qn('w:tblPr'))
    tw = etree.SubElement(tblPr, qn('w:tblW')); tw.set(qn('w:w'),'5289'); tw.set(qn('w:type'),'pct')
    for e in ['top','left','bottom','right','insideH','insideV']:
        b = etree.SubElement(tblPr, qn(f'w:{e}')); b.set(qn('w:val'),'single')
        b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'0'); b.set(qn('w:color'),'000000')
    look = etree.SubElement(tblPr, qn('w:tblLook')); look.set(qn('w:val'),'0000')

    tblGrid = etree.SubElement(tbl, qn('w:tblGrid'))
    for w in COL_W:
        etree.SubElement(tblGrid, qn('w:gridCol')).set(qn('w:w'),str(w))

    # 表頭兩列 (deepcopy from template)
    for tr in title_trs:
        tbl.append(deepcopy(tr))

    # 資料列
    offset = pair_idx
    for pi in range(n_pairs):
        odd_src, even_src = raw_pairs[pair_idx]
        seq = offset + pi + 1

        # 奇數列文字
        odd_text = [str(seq)] + [copy_text(odd_src, ci) for ci in range(1, 15)]
        if pi == 0:
            odd_text[3] = '否'
        else:
            odd_text[3] = ''

        tr1 = etree.SubElement(tbl, qn('w:tr'))
        trPr1 = etree.SubElement(tr1, qn('w:trPr'))
        th1 = etree.SubElement(trPr1, qn('w:trHeight'))
        th1.set(qn('w:val'),'226'); th1.set(qn('w:hRule'),'atLeast')
        for ci, txt in enumerate(odd_text):
            fsz = 22
            if ci == 4 and txt.replace('`', '') == '施作前14日':
                fsz = 18
            add_cell(tr1, txt, COL_W[ci], center=ci != 1, bold=False,
                     font_size=fsz,
                     merge_restart=True if ci in DAT_MERGE else None,
                     line_twip=LINE_H_TWIP)

        # 偶數列文字
        even_c1 = copy_text(even_src, 1)
        even_c4 = copy_text(even_src, 4)
        even_c5 = copy_text(even_src, 5)
        even_c13 = copy_text(even_src, 13)
        even_text = ['', even_c1, '', '', even_c4, even_c5, '',
                     '', '', '', '', '', '', even_c13, '']

        tr2 = etree.SubElement(tbl, qn('w:tr'))
        trPr2 = etree.SubElement(tr2, qn('w:trPr'))
        th2 = etree.SubElement(trPr2, qn('w:trHeight'))
        th2.set(qn('w:val'),'226'); th2.set(qn('w:hRule'),'atLeast')
        for ci, txt in enumerate(even_text):
            add_cell(tr2, txt, COL_W[ci], center=ci != 1, bold=False,
                     font_size=22,
                     merge_restart=False if ci in DAT_MERGE else None,
                     line_twip=LINE_H_TWIP)

        pair_idx += 1

body.append(etree.SubElement(body, qn('w:p')))
body.append(sect_pr)
tmpl.save(OUT)
print(f'Done! Pages={len(PAGE_PLAN)}, Items={pair_idx}')
print(f'File: {OUT}')
