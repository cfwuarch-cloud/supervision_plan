# -*- coding: utf-8 -*-
"""
列高校準工具 — 比對文字字數與實際 Word 渲染列高

使用方式：
  python -X utf8 tools/calibrate_row_height.py

輸出：
  - output/calibrate_row_height.docx  — 校準用 Word 文件
  - 終端機印出比對表格
"""
import sys, os, math, json, importlib.util
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from PIL import ImageFont, ImageDraw, Image

spec = importlib.util.spec_from_file_location('v2',
    os.path.join(os.path.dirname(os.path.dirname(__file__)),
                 'tables', 'table5.1', 'convert_5.1_v2.py'))
_mod = importlib.util.module_from_spec(spec)
spec.loader.exec_module(_mod)
load_price_sheet = _mod.load_price_sheet
C1_WIDTH_TWIP = _mod.C1_WIDTH_TWIP
clean_text = _mod.clean_text

PRICE_PATH = os.path.normpath(os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    'data', '02_成德-詳細價目表.xlsx'))
OUTPUT_PATH = os.path.normpath(os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    'output', 'calibrate_row_height.docx'))


def _get_font():
    try:
        return ImageFont.truetype('kaiu', 10)
    except:
        return ImageFont.load_default()


def calc_pillow_lines(text, col_width):
    font = _get_font()
    draw = ImageDraw.Draw(Image.new('RGB', (1, 1)))
    segments = text.split(chr(10))
    total = 0
    for seg in segments:
        if not seg.strip():
            total += 1
            continue
        bbox = draw.textbbox((0, 0), seg, font=font)
        w_twip = (bbox[2] - bbox[0]) * 20
        total += max(1, math.ceil(w_twip / col_width))
    return total


def main():
    items = load_price_sheet(PRICE_PATH, {'式', '工'})
    print(f'載入 {len(items)} 項')

    doc = Document()
    body = doc.element.body

    # 建立表格：1 欄 — 材料名稱
    tbl = etree.SubElement(body, qn('w:tbl'))
    tblPr = etree.SubElement(tbl, qn('w:tblPr'))
    tblW = etree.SubElement(tblPr, qn('w:tblW'))
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'dxa')
    tblGrid = etree.SubElement(tbl, qn('w:tblGrid'))
    gc = etree.SubElement(tblGrid, qn('w:gridCol'))
    gc.set(qn('w:w'), '5000')

    sample_count = min(50, len(items))
    results = []

    for i in range(sample_count):
        item, name, qty = items[i]
        lines_p = calc_pillow_lines(name, C1_WIDTH_TWIP)

        tr = etree.SubElement(tbl, qn('w:tr'))
        trPr = etree.SubElement(tr, qn('w:trPr'))

        # trHeight = atLeast 0 (no constraint, Word auto-calculates)
        th = etree.SubElement(trPr, qn('w:trHeight'))
        th.set(qn('w:val'), '0')
        th.set(qn('w:hRule'), 'atLeast')

        tc = etree.SubElement(tr, qn('w:tc'))
        tcPr = etree.SubElement(tc, qn('w:tcPr'))
        tcW = etree.SubElement(tcPr, qn('w:tcW'))
        tcW.set(qn('w:w'), '5000')
        tcW.set(qn('w:type'), 'dxa')

        p = etree.SubElement(tc, qn('w:p'))
        pPr = etree.SubElement(p, qn('w:pPr'))
        sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '0')

        r = etree.SubElement(p, qn('w:r'))
        rPr = etree.SubElement(r, qn('w:rPr'))
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:ascii'), '標楷體')
        rFonts.set(qn('w:eastAsia'), '標楷體')
        sz = etree.SubElement(rPr, qn('w:sz'))
        sz.set(qn('w:val'), '20')
        szCs = etree.SubElement(rPr, qn('w:szCs'))
        szCs.set(qn('w:val'), '20')
        t = etree.SubElement(r, qn('w:t'))
        t.text = name
        t.set(qn('xml:space'), 'preserve')

        results.append((i, len(name), lines_p, name[:30]))

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)

    # 讀回實際渲染列高
    doc2 = Document(OUTPUT_PATH)
    tbl2 = doc2.tables[0]
    print(f'\n{"#":>3} {"字數":>4} {"Pillow行":>7} {"實際高":>7} {"行高/行":>7}  名稱')
    print('-' * 60)
    for i, r in enumerate(results):
        if i >= len(tbl2.rows):
            break
        tr_xml = tbl2.rows[i]._tr
        trPr = tr_xml.find(qn('w:trPr'))
        actual_h = 0
        if trPr is not None:
            th = trPr.find(qn('w:trHeight'))
            if th is not None:
                actual_h = int(th.get(qn('w:val')))
        lines = r[2]
        per_line = round(actual_h / lines) if lines > 0 else 0
        print(f'{r[0]:>3} {r[1]:>4} {lines:>7} {actual_h:>7} {per_line:>7}  {r[3]}')


if __name__ == '__main__':
    main()
